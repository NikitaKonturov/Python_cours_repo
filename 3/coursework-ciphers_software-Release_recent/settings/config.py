import logging
import threading
import os
import re
import time
from pathlib import Path
from typing import BinaryIO

import uvicorn
import webview
from docx import Document
from dotenv import dotenv_values, find_dotenv, load_dotenv, set_key
from pydantic import Field, ValidationError
from pydantic_settings import BaseSettings
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response

# logger = logging.getLogger('Logger')
# logger.setLevel(logging.DEBUG)

# formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

# file_handler = logging.FileHandler('app.log')
# file_handler.setLevel(logging.DEBUG)
# file_handler.setFormatter(formatter)

# console_handler = logging.StreamHandler()
# console_handler.setLevel(logging.INFO)
# console_handler.setFormatter(formatter)

# logger.addHandler(file_handler)
# logger.addHandler(console_handler)


class NoCacheMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        response = await call_next(request)
        
        # Применяем только к статическим файлам и HTML
        path = request.url.path
        if (path.endswith(('.html', '.js', '.css', '.png', '.jpg', '.jpeg', '.gif')) or 
            path.startswith('/static/')):
            
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
            
            # Добавляем уникальный ETag на основе времени
            response.headers["ETag"] = f'"{int(time.time())}"'
        
        return response

# ================================== Docx/Txt Converters ==========================================


class FileLanguageError(Exception):
    def __init__(self, errorLanguage: str, message: str = "Неизвестный язык файла."):
        self.errorLanguage = errorLanguage
        self.message = message
        super().__init__(message)

    def __str__(self):
        return f"{self.message} (Язык: {self.errorLanguage})"

def match(text, alphabet=set('абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ')):
    return not alphabet.isdisjoint(text)

def check_file_path(filePath: str):
    if not os.path.exists(filePath):
        raise FileExistsError(f'Файл {filePath} не существует')


def save_open_text_docx_as_bin_file(language: str, openTextFile: BinaryIO, saveOpenTextTxtFile: Path):

    doc = Document(openTextFile)

    if (language.lower() == 'ru'):
        with open(saveOpenTextTxtFile, 'bw') as binFile:
            for para in doc.paragraphs:
                # checkPart = ''.join(re.findall(r'[A-Za-z]', para.text))
                cleanedText = ''.join(re.findall(
                    r'[А-Яа-я]', para.text)).upper()

                # if checkPart:
                #     raise FileLanguageError(
                #         f'Файл содержит символы другого языка!', errorLanguage="en")

                if cleanedText:
                    binFile.write(cleanedText.encode("utf-16-le"))

    elif (language.lower() == 'en'):
        with open(saveOpenTextTxtFile, 'bw') as binFile:
            for para in doc.paragraphs:
                # checkPart = ''.join(re.findall(r'[А-Яа-я]', para.text))
                cleanedText = ''.join(re.findall(
                    r'[A-Za-z]', para.text)).upper()

                # if checkPart:
                #     raise FileLanguageError(
                #         f'Файл содержит символы другого языка!', errorLanguage="ru")

                if cleanedText:
                    binFile.write(cleanedText.encode("utf-16-le"))

    else:
        raise FileLanguageError(errorLanguage="Anny", message='Неверный язык!')


def save_docx_as_txt(textFile: BinaryIO, saveTxtFile: Path):
    doc = Document(textFile)
    with open(saveTxtFile, "w", encoding='utf-8') as txtFile:
        for paragraph in doc.paragraphs:
            txtFile.write(paragraph.text)

    return


def save_open_text_as_bin_file(language: str, file: BinaryIO, pathToSaveTxt: Path, bufferSize: int = 1024):
    with open(pathToSaveTxt, "wb") as resBinFile:
        dataText : str = file.read().decode("utf-8")
        if (language.lower() == "ru"):
            cleanedText = ''.join(re.findall( r'[А-Яа-я]', dataText)).upper()
            print(cleanedText)
            if cleanedText:
                resBinFile.write(cleanedText.encode("utf-16-le"))
        elif (language.lower() == "en"):
            cleanedText = ''.join(re.findall(
                r'[A-Za-z]', dataText)).upper()
            if cleanedText:
                resBinFile.write(cleanedText.encode("utf-16-le"))
        else:
            raise FileLanguageError(
                f'The language is not defined. Supported languages: ru, en', errorLanguage="any")
    return


def save_as_txt_file(file: BinaryIO, pathToSaveTxtFile: Path, bufferSize: int = 20):
    with open(pathToSaveTxtFile, "w", encoding="utf-8") as resTxtFile:
        dataBuffer = file.read().decode("utf-8")
        resTxtFile.write(dataBuffer)
        print(dataBuffer)
    return


def check_file_path(filePath: str):
    if os.path.exists(filePath):
        print(f'The file {filePath} exists, continuing work...')
    else:
        raise Exception(f'The file {filePath} does not exist')





def check_path(path: Path) -> Path | None:
    if not path.exists():
        raise ValidationError(f"Path {str(path)} was not found...")
    return path


def check_host(host: str) -> str | None:
    if not re.match(r'[0-9]+\.[0-9]+\.[0-9]+\.[0-9]+', host):
        raise ValidationError(f"Invalid format for host: {host}")
    return host


# Путь к файлу .env
dotenv_path = find_dotenv("Config.env")

# Загружаем переменные окружения
env_values = dotenv_values(dotenv_path)


class Settings(BaseSettings):
    app_name: str = Field(..., env="APP_NAME")
    fiveGramsEnabled: str = Field(..., env="FIVEGRAMSENABLED")
    base_dir_path: Path = Field(..., env="BASE_DIR_PATH")
    path_to_ciphers: Path = Field(..., env="PATH_TO_CIPHERS")
    path_to_templates: Path = Field(..., env="PATH_TO_TEMPLATES")
    path_to_static: Path = Field(..., env="PATH_TO_STATIC")
    filename_to_save_full_open_text: str = Field(
        ..., env="FILENAME_TO_SAVE_FULL_OPEN_TEXT")
    encript_results_path: Path = Field(..., env="ENCRIPT_RESULTS_PATH")
    decript_results_path: Path = Field(..., env="DECRIPT_RESULTS_PATH")
    interface_language: str = Field(..., env="INTERFACE_LANGUAGE")
    ciphers_language: str = Field(..., env="CIPHERS_LANGUAGE")
    path_to_dir_viginer_dict: Path = Field(..., env="PATH_TO_DIR_VIGINER_DICT")
    host: str = Field(..., env="HOST")
    port: int = Field(..., env="PORT")
    location: str = Field(..., env="LOCATION")

    class Config:
        env_file = dotenv_path  # Указываем файл для поиска переменных окружения
        env_file_encoding = "utf-8"

    def update_settings(self, field: str, value: str) -> None:
        print(field, value)
        path_to_env = find_dotenv(dotenv_path)
        if not hasattr(self, field):
            raise ValueError(f"{field} was not found...")
        setattr(self, field, value)
        set_key(path_to_env, field.upper(), value)


def load_settings(config_filename: str) -> None:
    path_to_env = find_dotenv(config_filename)
    if not path_to_env:
        raise ValueError(f"{config_filename} was not found...")
    load_dotenv(path_to_env)


def update_js_file(pathToJsFile: Path, parametrs: dict[str, str]) -> None:
    with open(pathToJsFile, "r", encoding="utf-8") as file:
        code = file.read()
    print(parametrs)

    for key, value in parametrs.items():
        value = re.escape(str(value))
        pattern = rf"({key}\s*=\s*['\"])[^'\"]*(['\"];)"
        code = re.sub(pattern, rf'\1{value}\2', code)

    with open(pathToJsFile, "w", encoding="utf-8") as file:
        file.write(code)
        
def insert_every_n(text, symbol, n=6):
    return symbol.join(text[i:i+n] for i in range(0, len(text), n))
    


def save_to_docx(data: list[dict[str, str]], docxFile: Path, fiveGrams: str):
    check_file_path(docxFile.parent)
    doc = Document()
    for dataDict in data:
        for key in dataDict:
            doc.add_paragraph(key)
            text = dataDict[key]
            text = text.replace(' ', '')
            if(fiveGrams == "true"):
                text = insert_every_n(text, ' ', 5)
            doc.add_paragraph(text)
            doc.save(docxFile.__str__())
    return

EXCLUDED_DIRS = {
    "$Recycle.Bin",
    "$RECYCLE.BIN"
    "System Volume Information",
    "Windows",
    "Program Files",
    "Program Files (x86)",
    "ProgramData",
    "AppData",
    "$RECYCLE.BIN",
    "$RECYCLE.BIN",
    "SYSTEM VOLUME INFORMATION",
    "WINDOWS",
    "PROGRAM FILES",
    "PROGRAM FILES (X86)",
    "PROGRAMDATA",
    "APPDATA"
}


def search_directory(basePath: Path, dirname: str) -> None | Path:
    for root, dirs, _ in os.walk("C:\\"):
        # Фильтруем системные папки
        dirs[:] = [d for d in dirs if d not in EXCLUDED_DIRS]

        if dirname in dirs:
            return Path(root) / dirname  # Возвращаем полный путь
    return None


def start_server(settings: Settings, BASE_DIR: Path) -> None:
    update_js_file(Path(BASE_DIR, "static", "settingWindow.js"), {'interfaceLanguage': settings.interface_language, 'cipherLanguage': settings.ciphers_language, 'encryptFolderPath': settings.encript_results_path.name, 'decryptFolderPath': settings.decript_results_path.name})    
    uvicorn.run("__main__:app", host=settings.host,
                port=settings.port, reload=False)


def start_webview(settings: Settings) -> None:
    time.sleep(1)
    webview.create_window(settings.app_name, settings.location, width=1280, height=720)
    webview.start()
    
    
    
